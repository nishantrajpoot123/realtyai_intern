import requests
import json
import hashlib
import time
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# API Details
url_pre = "https://ap-east-1.tensorart.cloud"
api_key = "38dbb245-327b-485c-bfdf-63ae966edb73"  # Replace with your actual API key

# API Endpoints
url_job = "/v1/jobs"
url_resource = "/v1/resource"
url_workflow = "/v1/workflows"

# Headers for API Key Authentication
HEADERS = {
    'Content-Type': 'application/json',
    'Accept': 'application/json',
    'Authorization': f'Bearer {api_key}'
}

# Story scenes with titles and prompts - Modified for Arjun (male character)
STORY_SCENES = [
    # Scene 1
    {
        "title": "Arjun finds a mysterious, glowing box in his attic. Inside lie the time-traveling shoes.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in a cozy attic, wearing pajamas and white shoes, surrounded by dusty books and glowing magical shoes, warm and inviting atmosphere, storybook style, watercolor style"
    },
    
    # Scene 2
    {
        "title": "Arjun slips the shoes on for the first time and there feels some magic inside the room.",
        "prompt": "Arjun, a 5 years old fair and white complexion kid, wears the magical white shoes and the room is filled with magical vibes. Visuals: A warm, inviting room with an aura of mystery, soft magical light around, kids storybook style, watercolor style"
    },
    
    # Scene 3
    {
        "title": "Transported to a lush, prehistoric jungle, Arjun encounters friendly, cartoonish dinosaurs.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion explores a prehistoric jungle, wearing a simple adventure outfit and white shoes. He walks among giant ferns, surrounded by towering trees and a distant volcano. The jungle is vibrant, with oversized leaves and a warm, ancient atmosphere. Kids storybook style, watercolor style."
    },
    
    # Scene 4
    {
        "title": "Arjun steps into the sands of ancient Egypt, where he marvels at towering pyramids and meets a wise camel near a pond.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in ancient Egypt, dressed as a little prince with a golden headdress, wearing white shoes, standing beside a majestic pyramid, helping a lost cat, warm desert tones, storybook style, watercolor style."
    },
    
    # Scene 5
    {
        "title": "In a bustling Roman marketplace, Arjun is dressed in a simple tunic.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in ancient Rome, wearing a simple tunic, a laurel wreath, and white shoes, exploring the grand Colosseum, helping a boy carry water, historical and grand, storybook style, watercolor style"
    },
    
    # Scene 6
    {
        "title": "Arjun finds himself in a medieval village near a grand castle.",
        "prompt": "Arjun, a 5 years old fair and white complexion kid, walks in a medieval village near a grand castle. Visuals: Cobblestone streets, rustic cottages, and a mystical forest with shimmering lights filtering through ancient trees, kids storybook style, watercolor style"
    },
    
    # Scene 7
    {
        "title": "Now in the heart of the Renaissance, Arjun is transformed into a budding inventor.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in a Renaissance-era city, wearing an ornate inventor's outfit, surrounded by sketches of inventions, wooden tables with gadgets, and colorful market scenes. He stands in a vibrant workshop filled with tools and parchment, with an art-filled piazza in the background. Storybook style, watercolor style."
    },
    
    # Scene 8
    {
        "title": "In a foggy, gaslit Victorian city, Arjun sports a smart outfit complete with a miniature top hat.",
        "prompt": "In a foggy, gaslit Victorian city, Arjun, a 5 years old fair and white complexion kid, sports a smart outfit complete with a miniature top hat. He explores secret alleys and marvels at fantastical steam-powered contraptions. Visuals: Moody, atmospheric streets, delicate gaslight reflections, and whimsical steampunk elements. kids storybook style, watercolor style, full body."
    },
    
    # Scene 9
    {
        "title": "Arjun rides through a dusty frontier town on horseback.",
        "prompt": "Arjun, a 5 years old fair and white complexion kid, rides through a dusty frontier town on horseback, dressed in a rugged, yet playful Western outfit. A friendly, animated tumbleweed and a wise old cactus share clues to his next destination. Visuals: Sunlit desert landscapes, vintage saloon facades, and caricatured cowboy details that keep the tone light."
    },
    
    # Scene 10
    {
        "title": "Arjun steps into a lively 1920s street scene, with jazz music filling the air.",
        "prompt": "Arjun, a 5 years old fair and white complexion kid, lives in 1920s street scene, with jazz music filling the air. His outfit is chic and era-appropriate, and a vintage car and dancing silhouettes add flair to the backdrop. Visuals: Bold art deco designs, a lively dance scene under streetlamps, and dynamic expressions of joy and movement. kids storybook style, watercolor style"
    },
    
    # Scene 11
    {
        "title": "Arjun stands in city of flying cars and holographic billboards, befriending a quirky robot guide.",
        "prompt": "In a dazzling futuristic metropolis, Arjun, a 5 years old fair and white complexion kid, wears a sleek, modern outfit accented with neon. He navigates a city of flying cars and holographic billboards, befriending a quirky robot guide. Visuals: A high-tech skyline with luminous colors, dynamic digital displays, and a harmonious blend of technology and art. kids storybook style, watercolor style"
    },
    
    # Scene 12
    {
        "title": "Arjun enters a surreal digital world where landscapes morph with every step.",
        "prompt": "Arjun, a 5 years old fair and white complexion kid, enters a surreal digital world where landscapes morph with every step. Here, the scenery is pixelated, and his outfit shifts with vibrant, glitch-like effects as he interacts with fantastical virtual creatures. Visuals: A playful mix of neon pixels and digital motifs that create a dreamlike, abstract adventure. kids storybook style, watercolor style"
    },
    
    # Scene 13
    {
        "title": "Arjun now jumps into the world of Flowers in a silent old pleasant village.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion stands in the home garden in a village under a serene, starlit sky. His softly glowing shoes shimmer like distant memories. He is surrounded by blooming flowers and rustling leaves, bathed in gentle starlight, with an aura of mystery and wisdom. Storybook style, watercolor style."
    },
    
    # Scene 14
    {
        "title": "In ancient Rome, Arjun walks through the streets near the grand arena.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in ancient Rome, wearing a simple tunic, a laurel wreath, and white shoes, exploring the grand Colosseum, helping a boy carry water, historical and grand, storybook style, watercolor style."
    },
    
    # Scene 15
    {
        "title": "Arjun visits a magnificent Persian palace and learns the magic.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in a Persian palace, wearing an elegant embroidered robe, golden accessories, and white shoes, surrounded by glowing lanterns and intricate tilework, magical and rich, storybook style, watercolor style."
    },
    
    # Scene 16
    {
        "title": "In Victorian London, Arjun explores an old bookshop filled with magical stories.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in Victorian London, wearing a frilly outfit, a cap, and white shoes, exploring an old bookshop with gas lanterns glowing outside, cozy and vintage, storybook style, watercolor style."
    },
    
    # Scene 17
    {
        "title": "Under the cherry blossom trees, Arjun learns the art of origami from a kind samurai.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in ancient Japan, wearing a simple kimono with traditional patterns, white shoes, standing near a river with floating paper lanterns, peaceful and serene, storybook style, watercolor style."
    },
    
    # Scene 18
    {
        "title": "Arjun listens to thrilling stories about dragons and sea voyages around a warm Viking campfire.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in a Viking village, wearing a warm fur cloak and white shoes, standing near a wooden longship, listening to an elder tell stories by the fire, adventurous and rustic, storybook style, watercolor style."
    },
    
    # Scene 19
    {
        "title": "In the Mayan jungle, Arjun helps a young explorer decode ancient glowing symbols.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in an ancient Mayan jungle, wearing a woven tunic, colorful feathers in his hair, and white shoes, surrounded by glowing temple glyphs and lush greenery, mysterious and exotic symbols on the trees, storybook style, watercolor style"
    },
    
    # Scene 20
    {
        "title": "Arjun flies a red paper kite high above the Great Wall with his new friend.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion on the Great Wall of China, wearing a traditional silk outfit with dragon embroidery, white shoes, flying a red paper kite with a child, breathtaking and historical, storybook style, watercolor style."
    },
    
    # Scene 21
    {
        "title": "In the Wild West, Arjun rides a pony across golden plains.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in the Wild West, wearing a cowboy hat, denim outfit, and white shoes, riding a pony across golden plains, adventurous and rustic, storybook style, watercolor style."
    },
    
    # Scene 22
    {
        "title": "Arjun learns how to weave colorful blankets with the villagers high in the mountains.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in an Incan mountain village, wearing a vibrant poncho, white shoes, helping villagers weave colorful textiles, surrounded by llamas and ancient stone ruins, peaceful and cultural, storybook style, watercolor style."
    },
    
    # Scene 23
    {
        "title": "Arjun helps collect water from the river while watching giraffes and zebras roam free.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in an African safari village, wearing a tribal outfit with bold patterns, white shoes, standing among zebras and giraffes, helping an elder collect water, wild and adventurous, storybook style, watercolor style."
    },
    
    # Scene 24
    {
        "title": "In a quiet Mongolian village, Arjun helps herd baby goats near a cozy round yurt.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in a Mongolian nomadic camp, wearing a thick fur-lined coat, white shoes, helping a child herd baby goats near a round yurt, vast and peaceful, storybook style, watercolor style."
    },
    
    # Scene 25
    {
        "title": "Arjun watches an exciting play under the sunset in a grand Greek amphitheater.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in an ancient Greek amphitheater, wearing a white tunic with golden trim, white shoes, watching a grand play under a sunset sky, classical and dramatic, storybook style, watercolor style."
    },
    
    # Scene 26
    {
        "title": "Deep underwater, Arjun explores the glowing ruins of the legendary lost city of Atlantis.",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in Atlantis, wearing a shimmering blue outfit, white shoes, surrounded by glowing fish and coral towers, exploring the underwater city, mystical and dreamlike, storybook style, watercolor style."
    },
    
    # Scene 27
    {
        "title": "Arjun wakes up back in his attic, realizing his magic shoes have taken him on an incredible journey!",
        "prompt": "Arjun, a young 5-year-old boy with a fair and white complexion in his cozy attic, wearing pajamas and white shoes, looking at the magical shoes glowing softly, surrounded by books and childhood toys, warm and nostalgic, storybook style, watercolor style."
    }
]

def ensure_output_folders():
    """Ensure that the 'outputs' folder exists."""
    if not os.path.exists("outputs"):
        os.makedirs("outputs")
    if not os.path.exists("outputs/images"):
        os.makedirs("outputs/images")

def get_job_result(job_id, scene_number):
    """Poll the API for job completion status and save the resulting image."""
    while True:
        time.sleep(5)  # Polling interval
        response = requests.get(f"{url_pre}{url_job}/{job_id}", headers=HEADERS)
        
        if response.status_code != 200:
            print(f"Error checking job status: {response.status_code}")
            time.sleep(10)
            continue
            
        job_response = response.json()
        if 'job' in job_response:
            job_dict = job_response['job']
            job_status = job_dict.get('status')
            print(f"Scene {scene_number} - Job status: {job_status}")
            
            if job_status == 'SUCCESS':
                image_url = job_dict['successInfo']['images'][0]['url']
                print(f"Scene {scene_number} - Image URL: {image_url}")
                image_path = save_image(image_url, scene_number)
                return image_path
            elif job_status == 'FAILED':
                print(f"Scene {scene_number} - Image generation failed.")
                print(job_dict.get('failedInfo', {}))
                return None
            elif job_status == 'PROCESSING':
                print(f"Scene {scene_number} - Still processing... waiting.")
            else:
                print(f"Scene {scene_number} - Status: {job_status}")
        else:
            print(f"Scene {scene_number} - Unexpected response: {job_response}")
        
        time.sleep(10)  # Wait before checking again

def save_image(image_url, scene_number):
    """Download and save the generated image with scene number."""
    ensure_output_folders()
    response = requests.get(image_url)
    if response.status_code == 200:
        # Create a filename with scene number for better organization
        image_path = os.path.join("outputs/images", f"scene_{scene_number:02d}_arjun.png")
        with open(image_path, "wb") as img_file:
            img_file.write(response.content)
        print(f"Scene {scene_number} - Image saved to {image_path}")
        return image_path
    else:
        print(f"Scene {scene_number} - Error downloading image: {response.status_code}")
        return None

def upload_img(img_path):
    """Upload an image and return the resource ID."""
    print(f"Uploading image: {img_path}")
    data = {"expireSec": 3600}
    response = requests.post(f"{url_pre}{url_resource}/image", json=data, headers=HEADERS)
    print("Upload response status:", response.status_code)
    
    if response.status_code != 200:
        print(f"Error uploading image: {response.text}")
        return None
    
    response_data = response.json()
    resource_id = response_data['resourceId']
    put_url = response_data['putUrl']
    headers = response_data['headers']
    
    with open(img_path, 'rb') as f:
        res = f.read()
        upload_response = requests.put(put_url, data=res, headers=headers)
        print("PUT response status:", upload_response.status_code)
        
        if upload_response.status_code >= 300:
            print(f"Error in PUT request: {upload_response.text}")
            return None
    
    return resource_id

def generate_scene_image(img_path, scene_number, scene_data):
    """Generate an image for a specific scene."""
    print(f"\n=== Processing Scene {scene_number}: {scene_data['title']} ===")
    
    # Upload the reference image
    resource_id = upload_img(img_path)
    if not resource_id:
        print(f"Failed to upload image for scene {scene_number}")
        return None
    
    # Build the payload with all required field attributes
    data = {
        "request_id": hashlib.md5(f"{int(time.time())}_{scene_number}".encode()).hexdigest(),
        "templateId": "688362427502551075",  # Template URL: https://tensor.art/template/688362427502551075
        "fields": {
            "fieldAttrs": [
                # Node 12: LoadImage (FACE)
                {"nodeId": "12", "fieldName": "image", "fieldValue": resource_id},
                # Node 25: SDXL Prompt Styler
                {"nodeId": "25", "fieldName": "style", "fieldValue": "sai-line art"},
                {"nodeId": "25", "fieldName": "text_positive", "fieldValue": scene_data['prompt']},
                {"nodeId": "25", "fieldName": "text_negative", "fieldValue": 
                 "lowres, bad hands, text, error, missing fingers, extra digit, fewer digits, "
                 "cropped, worst quality, low quality, normal quality, jpeg artifacts, signature, "
                 "watermark, username, blurry, patreon logo, artist name, sexual content, adult, nuidity"},
                # Node 18: IPAdapter FaceID
                {"nodeId": "18", "fieldName": "weight", "fieldValue": "0.8"},
                # Node 3: KSampler
                {"nodeId": "3", "fieldName": "seed", "fieldValue": str(161098661698898 + scene_number)},  # Unique seed per scene
                {"nodeId": "3", "fieldName": "cfg", "fieldValue": "2"},
                {"nodeId": "3", "fieldName": "steps", "fieldValue": "20"},
                {"nodeId": "3", "fieldName": "sampler_name", "fieldValue": "dpmpp_sde"},
                {"nodeId": "3", "fieldName": "scheduler", "fieldValue": "karras"},
                # Node 5: Empty Latent Image
                {"nodeId": "5", "fieldName": "width", "fieldValue": "1024"},
                {"nodeId": "5", "fieldName": "height", "fieldValue": "1024"},
                # Node 4: Load Checkpoint
                {"nodeId": "4", "fieldName": "ckpt_name", "fieldValue": "676746065682318967"},
                # Node 38: LoraLoaderModelOnly
                {"nodeId": "38", "fieldName": "lora_name", "fieldValue": "681174344216573550"},
                {"nodeId": "38", "fieldName": "strength_model", "fieldValue": "0.6"},
                {"nodeId": "38", "fieldName": "tams_lora_name", "fieldValue": "IP_adapter_face_id -  SDXL"}
            ]
        },
    }
    
    # Submit the job using the workflow template endpoint
    print(f"Submitting job for scene {scene_number}")
    response = requests.post(f"{url_pre}{url_job}/workflow/template", json=data, headers=HEADERS)
    
    if response.status_code != 200:
        print(f"Error submitting job for scene {scene_number}: {response.text}")
        return None
    
    response_data = response.json()
    
    if 'job' in response_data:
        job_dict = response_data['job']
        job_id = job_dict.get('id')
        job_status = job_dict.get('status')
        print(f"Scene {scene_number} - Job ID: {job_id}, Initial Status: {job_status}")
        return get_job_result(job_id, scene_number)
    else:
        print(f"Failed to create workflow job for scene {scene_number}:", response_data)
        return None

def create_story_document(image_paths):
    """Create a Word document with scenes, prompts, and images."""
    doc = Document()
    
    # Add title
    title = doc.add_heading("Arjun's Time-Traveling Shoes", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph("A magical journey through time")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add premise
    doc.add_heading("Premise:", 1)
    premise = doc.add_paragraph("Arjun, a 5 years old kid, discovers a magical pair of shoes that transport him across time. In each era, he's dressed in a style that fits the period, meets fascinating creatures or nonhuman characters, and experiences adventures that blend history with fantasy. The story unfolds through immersive scenes.")
    
    # Line break
    doc.add_paragraph("")
    
    # Add each scene
    for i, scene_data in enumerate(STORY_SCENES, 1):
        # Add scene heading
        scene_heading = doc.add_heading(f"Scene {i}: {scene_data['title']}", 2)
        
        # Add prompt subheading
        doc.add_heading("Prompt:", 3)
        
        # Add prompt content
        prompt_para = doc.add_paragraph(scene_data['prompt'])
        
        # Add image if available
        if i in image_paths and image_paths[i]:
            try:
                doc.add_picture(image_paths[i], width=Inches(6))
                image_para = doc.paragraphs[-1]
                image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error adding image for scene {i}: {e}")
        
        # Add some space between scenes
        doc.add_paragraph("")
    
    # Save the document
    doc_path = os.path.join("outputs", "Arjuns_Time_Traveling_Shoes.docx")
    doc.save(doc_path)
    print(f"\nStory document created: {doc_path}")
    return doc_path

def process_all_scenes(img_path):
    """Process all scenes using the provided reference image."""
    ensure_output_folders()
    
    successful_scenes = 0
    failed_scenes = []
    image_paths = {}
    
    for i, scene_data in enumerate(STORY_SCENES, 1):
        print(f"\n==================================================")
        print(f"GENERATING SCENE {i} OF {len(STORY_SCENES)}")
        print(f"==================================================")
        
        image_path = generate_scene_image(img_path, i, scene_data)
        
        if image_path:
            successful_scenes += 1
            image_paths[i] = image_path
        else:
            failed_scenes.append(i)
        
        # Add a delay between API calls to avoid rate limiting
        if i < len(STORY_SCENES):
            wait_time = 5
            print(f"Waiting {wait_time} seconds before processing next scene...")
            time.sleep(wait_time)
    
    # Print summary
    print("\n==================================================")
    print(f"GENERATION COMPLETE: {successful_scenes} of {len(STORY_SCENES)} scenes generated successfully")
    if failed_scenes:
        print(f"Failed scenes: {failed_scenes}")
    print("==================================================")
    
    # Create the story document
    doc_path = create_story_document(image_paths)
    print(f"Document created at: {doc_path}")
    
    return successful_scenes, failed_scenes, doc_path

if __name__ == '__main__':
    print("ARJUN'S TIME-TRAVELING SHOES - IMAGE GENERATOR")
    print("==============================================")
    print("This program will generate images for all scenes in the story")
    print("and create a Word document with the complete story.\n")
    
    # Check if python-docx is installed
    try:
        from docx import Document
    except ImportError:
        print("ERROR: The 'python-docx' library is not installed.")
        print("Please install it by running: pip install python-docx")
        exit(1)
    
    # Ask for the reference male character image
    ref_img_path = input("Enter the path to your reference male character image: ")
    
    # Validate that the file exists
    if not os.path.exists(ref_img_path):
        print(f"Error: File '{ref_img_path}' does not exist!")
    else:
        # Confirm with the user before starting
        print(f"\nAbout to generate 27 images for 'Arjun's Time-Traveling Shoes' story")
        print(f"Using reference image: {ref_img_path}")
        confirm = input("This will make multiple API calls. Continue? (y/n): ")
        
        if confirm.lower() == 'y':
            successful, failed, doc_path = process_all_scenes(ref_img_path)
            print(f"\nSUMMARY:")
            print(f"- {successful} scenes generated successfully")
            print(f"- {len(failed)} scenes failed")
            print(f"- Document created: {doc_path}")
        else:
            print("Operation cancelled by user.")